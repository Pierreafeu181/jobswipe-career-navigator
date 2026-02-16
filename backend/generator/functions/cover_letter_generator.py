"""
cover_letter_gemini_docx_generator.py

Generate a FRENCH cover letter (lettre de motivation) as a DOCX,
personalized by Gemini, then convert it to PDF.

Inputs:
- offer_parsed: output of parse_job_offer_gemini(offer_text)
- cv_parsed:    output of parse_cv_with_gemini(cv_text)

Gemini:
- Generates the letter text in French, following a structure similar to:

  Prénom et Nom
  Adresse
  Mail et numéro de téléphone

  Contact dans l’entreprise (si vous l’avez)
  Nom de l’entreprise
  Adresse

  Fait à [Ville], le [Date]
  Objet : Candidature pour le poste de [Nom du poste] - [Référence]
  Madame, Monsieur,
  [Paragraphe 1]
  [Paragraphe 2]
  [Paragraphe 3]
  [Paragraphe 4]
  [Nom et Prénom]

Dependencies:
    pip install google-generativeai python-docx docx2pdf
"""

from __future__ import annotations

from typing import Dict, Any, Optional
import os
import json
import re
import datetime

from dotenv import load_dotenv
from google import genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pyhere import here

load_dotenv()

# ============================================================================
# 1. CONFIG GEMINI
# ============================================================================

try:
    from xhtml2pdf import pisa
except ImportError:
    raise ImportError("La librairie 'xhtml2pdf' est manquante. Installez-la avec : pip install xhtml2pdf")


# ============================================================================
# 2. Helpers (dates / normalisation / JSON)
# ============================================================================

def normalize_str(s: Optional[str]) -> str:
    if not s:
        return ""
    return " ".join(s.split())


def french_date(today: Optional[datetime.date] = None) -> str:
    if today is None:
        today = datetime.date.today()
    months = [
        "janvier", "février", "mars", "avril", "mai", "juin",
        "juillet", "août", "septembre", "octobre", "novembre", "décembre",
    ]
    return f"{today.day} {months[today.month - 1]} {today.year}"


def extract_json_from_output(output: str) -> Dict[str, Any]:
    """
    Extract a JSON object from the model's raw text output.

    - If the output is already pure JSON, parse directly.
    - Otherwise, find the first {...} block.
    """
    output = output.strip()

    # 1. Gestion des blocs Markdown (```json ... ```)
    if "```" in output:
        match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", output, flags=re.DOTALL)
        if match:
            output = match.group(1)
        else:
            output = re.sub(r"^```[a-zA-Z0-9]*\s*", "", output)
            output = re.sub(r"\s*```$", "", output)

    # 2. Extraction du bloc JSON
    match = re.search(r"\{.*\}", output, flags=re.DOTALL)
    if match:
        json_str = match.group(0)
    start = output.find("{")
    end = output.rfind("}")
    
    if start != -1 and end != -1 and end > start:
        json_str = output[start : end + 1]
    else:
        json_str = output

    # 3. Parsing avec tentative de nettoyage
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        # Tentative de réparation du JSON
        # 1. Suppression des commentaires // (sauf si dans une URL http://)
        json_str = re.sub(r"(?<!:)\/\/.*", "", json_str)
        # 2. Suppression des virgules traînantes (trailing commas)
        json_str = re.sub(r",\s*([\]}])", r"\1", json_str)
        # 3. Ajout des virgules manquantes entre un bloc fermant et une clé suivante
        json_str = re.sub(r"([\}\]])\s*(\"[a-zA-Z0-9_]+\"\s*:)", r"\1,\2", json_str)
        json_str = re.sub(r"([\}\]])\s*(\"[^\"]+\"\s*:)", r"\1,\2", json_str)
        # 4. Ajout des virgules manquantes après une valeur simple et une clé
        json_str = re.sub(r"([0-9]+|true|false|null)\s+(\"[a-zA-Z0-9_]+\"\s*:)", r"\1,\2", json_str)
        json_str = re.sub(r"([0-9]+|true|false|null)\s+(\"[^\"]+\"\s*:)", r"\1,\2", json_str)
        # 5. Ajout des virgules manquantes après une string et une clé
        json_str = re.sub(r"(\")\s+(\"[a-zA-Z0-9_]+\"\s*:)", r"\1,\2", json_str)
        json_str = re.sub(r"(\")\s+(\"[^\"]+\"\s*:)", r"\1,\2", json_str)

        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            raise ValueError(
                f"Erreur de parsing JSON : {e}\n"
                f"Sortie brute : {output[:500]}..."
            )


# ============================================================================
# 3. Extraction minimale des infos de contact pour aider Gemini
# ============================================================================

def pick_contact_info(cv_parsed: Dict[str, Any]) -> Dict[str, str]:
    first_name = normalize_str(cv_parsed.get("first_name"))
    last_name = normalize_str(cv_parsed.get("last_name"))
    full_name = normalize_str(cv_parsed.get("full_name"))

    if not full_name and (first_name or last_name):
        full_name = (first_name + " " + last_name).strip()

    contacts = cv_parsed.get("contacts") or {}
    emails = contacts.get("emails") or []
    phones = contacts.get("phones") or []
    locations = contacts.get("locations") or []

    email = emails[0] if emails else ""
    phone = phones[0] if phones else ""
    city = locations[0] if locations else ""

    return {
        "first_name": first_name,
        "last_name": last_name,
        "full_name": full_name,
        "email": email,
        "phone": phone,
        "city": city,
    }


# ============================================================================
# 4. Prompt Gemini pour générer la lettre
# ============================================================================
def build_header_prompt(
    offer_parsed: Dict[str, Any],
    cv_parsed: Dict[str, Any],
    city_hint: str,
    date_hint: str,
    reference: Optional[str] = None,
) -> str:
    offer_json = json.dumps(offer_parsed, ensure_ascii=False, indent=2)
    cv_json = json.dumps(cv_parsed, ensure_ascii=False, indent=2)

    return f"""
You are an expert French administrative assistant.
Your task is to prepare the HEADER and METADATA for a professional cover letter.

You MUST return STRICTLY a JSON object with this exact structure:
{{
  "header_blocks": {{
    "fullname_block": "Candidate Full Name",
    "location_block": "Candidate Address/City",
    "email_block": "Candidate Email",
    "phone_block": "Candidate Phone",
    "websites_block": "Candidate Links (LinkedIn, Portfolio...)"
  }},
  "company_blocks": {{
    "contact_block": "Recruiter Name (if known) or empty",
    "company_name_block": "Company Name",
    "company_address_block": "Company Address (if known) or City"
  }},
  "place_date_line": "Fait à [City], le [Date]",
  "objet_line": "Objet : Candidature pour le poste de [Job Title] [Reference]"
}}

Instructions:
- Use the candidate's real information from the CV JSON.
- Use the company's info from the Job Offer.
- Format 'place_date_line' using the city: "{city_hint}" and date: "{date_hint}".
- If a reference is provided ("{reference if reference else ''}"), include it in the object line.

JOB OFFER:
{offer_json}

CANDIDATE CV:
{cv_json}
"""


def build_body_prompt(
    offer_parsed: Dict[str, Any],
    cv_parsed: Dict[str, Any],
    gender_label: str,
    objet_line: str
) -> str:
    offer_json = json.dumps(offer_parsed, ensure_ascii=False, indent=2)
    cv_json = json.dumps(cv_parsed, ensure_ascii=False, indent=2)

    return f"""
You are an expert French copywriter specialized in cover letters.
Your task is to write the BODY of the letter.

Context:
- Candidate Gender: {gender_label}
- Letter Object: "{objet_line}"

You MUST return STRICTLY a JSON object with this exact structure:
{{
  "greeting": "Madame, Monsieur,",
  "para1": "Introduction paragraph (Hook)...",
  "para2": "Why me (Skills & Experience)...",
  "para3": "Why you (Company alignment)...",
  "para4": "Call to action (Interview request)...",
  "signature": "Candidate Name"
}}

Instructions:
- Write in professional French.
- "para1" MUST NOT contain the greeting.
- Adapt to the specific job offer (missions, technologies, skills).
- Use the candidate's real information from the CV.
- STRICT FIDELITY: Do NOT invent any experience or skill. Use ONLY what is in the CV.
- Be convincing but factual.

JOB OFFER:
{offer_json}

CANDIDATE CV:
{cv_json}
"""

def generate_letter_structure_with_gemini(
    offer_parsed: Dict[str, Any],
    cv_parsed: Dict[str, Any],
    gender: str = "M",
    reference: Optional[str] = None,
    city_override: Optional[str] = None,
    date_override: Optional[str] = None,
    api_key: str = "",
    model_name: str = "gemini-1.5-flash"
) -> Dict[str, Any]:
    # 1. Préparation des indices
    contact = pick_contact_info(cv_parsed)
    city_hint = city_override or contact["city"] or "Paris"
    date_hint = date_override or french_date()
    gender_label = "masculin" if gender.upper() == "M" else "féminin"
    
    client = genai.Client(api_key=api_key)

    # 2. Appel 1 : Header & Meta
    prompt_header = build_header_prompt(offer_parsed, cv_parsed, city_hint, date_hint, reference)
    resp_header = client.models.generate_content(
        model=model_name,
        contents=prompt_header
    )
    json_header = extract_json_from_output(resp_header.text)

    # Récupération de l'objet pour le contexte du body
    objet_line = json_header.get("objet_line", "")

    # 3. Appel 2 : Body
    prompt_body = build_body_prompt(offer_parsed, cv_parsed, gender_label, objet_line)
    resp_body = client.models.generate_content(
        model=model_name,
        contents=prompt_body
    )
    json_body = extract_json_from_output(resp_body.text)

    # 4. Fusion
    full_json = {**json_header, **json_body}

    # Helpers pour sécuriser les champs
    def norm(s: Any) -> str:
        if not isinstance(s, str):
            s = "" if s is None else str(s)
        return normalize_str(s)

    header_blocks_raw = full_json.get("header_blocks") or {}
    company_blocks_raw = full_json.get("company_blocks") or {}

    header_blocks = {
        "fullname_block": norm(header_blocks_raw.get("fullname_block", "")),
        "location_block": norm(header_blocks_raw.get("location_block", "")),
        "email_block": norm(header_blocks_raw.get("email_block", "")),
        "phone_block": norm(header_blocks_raw.get("phone_block", "")),
        "websites_block": norm(header_blocks_raw.get("websites_block", "")),
    }

    company_blocks = {
        "contact_block": norm(company_blocks_raw.get("contact_block", "")),
        "company_name_block": norm(company_blocks_raw.get("company_name_block", "")),
        "company_address_block": norm(company_blocks_raw.get("company_address_block", "")),
    }

    return {
        "header_blocks": header_blocks,
        "company_blocks": company_blocks,
        "place_date_line": norm(full_json.get("place_date_line", "")),
        "objet_line": norm(full_json.get("objet_line", "")),
        "greeting": norm(full_json.get("greeting", "")),
        "para1": norm(full_json.get("para1", "")),
        "para2": norm(full_json.get("para2", "")),
        "para3": norm(full_json.get("para3", "")),
        "para4": norm(full_json.get("para4", "")),
        "signature": norm(full_json.get("signature", "")),
    }


# ============================================================================
# 5b. Génération HTML pour PDF (Alternative robuste sans Word)
# ============================================================================

def build_cover_letter_html_from_chunks(chunks: Dict[str, Any]) -> str:
    """
    Construit le HTML de la lettre pour conversion PDF via xhtml2pdf.
    """
    def safe(val: Optional[str]) -> str:
        return (val or "").replace("\n", "<br/>")

    header = chunks.get("header_blocks", {}) or {}
    company = chunks.get("company_blocks", {}) or {}
    header = chunks.get("header_blocks") or {}
    company = chunks.get("company_blocks") or {}

    css = """
    @page { size: a4 portrait; margin: 2.5cm; }
    body { font-family: Helvetica, sans-serif; font-size: 11pt; line-height: 1.4; color: #000; }
    .sender { text-align: left; margin-bottom: 30px; font-size: 10pt; }
    .recipient { text-align: right; margin-bottom: 40px; font-size: 10pt; }
    .meta { text-align: right; margin-bottom: 30px; }
    .object { font-weight: bold; margin-bottom: 20px; }
    .content p { margin-bottom: 12px; text-align: justify; }
    .signature { margin-top: 40px; }
    """

    html = f"<html><head><style>{css}</style></head><body>"

    # Expéditeur
    html += "<div class='sender'>"
    if header.get("fullname_block"): html += f"<strong>{safe(header['fullname_block'])}</strong><br/>"
    if header.get("location_block"): html += f"{safe(header['location_block'])}<br/>"
    if header.get("email_block"): html += f"{safe(header['email_block'])}<br/>"
    if header.get("phone_block"): html += f"{safe(header['phone_block'])}<br/>"
    if header.get("websites_block"): html += f"{safe(header['websites_block'])}<br/>"
    html += "</div>"

    # Destinataire
    html += "<div class='recipient'>"
    if company.get("contact_block"): html += f"{safe(company['contact_block'])}<br/>"
    if company.get("company_name_block"): html += f"{safe(company['company_name_block'])}<br/>"
    if company.get("company_address_block"): html += f"{safe(company['company_address_block'])}<br/>"
    html += "</div>"

    # Date
    if chunks.get("place_date_line"):
        html += f"<div class='meta'>{safe(chunks['place_date_line'])}</div>"

    # Objet
    if chunks.get("objet_line"):
        html += f"<div class='object'>{safe(chunks['objet_line'])}</div>"

    # Corps
    html += "<div class='content'>"
    if chunks.get("greeting"): html += f"<p>{safe(chunks['greeting'])}</p>"
    for k in ["para1", "para2", "para3", "para4"]:
        if chunks.get(k): html += f"<p>{safe(chunks[k])}</p>"
    if chunks.get("signature"): html += f"<div class='signature'>{safe(chunks['signature'])}</div>"
    html += "</div></body></html>"

    return html

# ============================================================================
# 5. Génération DOCX à partir de la structure (sans logique métier)
# ============================================================================
def build_cover_letter_docx_from_chunks(
    chunks: Dict[str, str],
    output_path: str,
) -> str:
    """
    Écrit un DOCX de lettre de motivation à partir des chunks générés par Gemini.
    Version avec espaces réduits entre les blocs.
    """
    doc = Document()

    # Style global
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ==========================
    # HEADER CANDIDAT — plusieurs blocs alignés à gauche
    # ==========================
    header = chunks.get("header_blocks", {}) or {}

    def safe_block(val: str) -> Optional[str]:
        val = (val or "").strip()
        return val if val else None

    for key in ["fullname_block", "location_block", "email_block", "phone_block", "websites_block"]:
        line = safe_block(header.get(key))
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        if key == "fullname_block":
            run.bold = True

    # ==========================
    # COMPANY BLOCK — aligné à droite, blocs séparés
    # ==========================
    company = chunks.get("company_blocks", {}) or {}

    for key in ["contact_block", "company_name_block", "company_address_block"]:
        line = safe_block(company.get(key))
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(line)

    # ==========================
    # LIEU + DATE
    # ==========================
    place_date = safe_block(chunks.get("place_date_line"))
    if place_date:
        pd = doc.add_paragraph(place_date)
        pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ==========================
    # OBJET (gras)
    # ==========================
    objet_line = safe_block(chunks.get("objet_line"))
    if objet_line:
        p = doc.add_paragraph()
        run = p.add_run(objet_line)
        run.bold = True

    # ==========================
    # SALUTATION (greeting) avant para1
    # ==========================
    greet = safe_block(chunks.get("greeting"))
    if greet:
        doc.add_paragraph(greet)

    # pas de ligne blanche supplémentaire (demande utilisateur)
    # directement les paragraphes

    # ==========================
    # PARAGRAPHES 1 à 4
    # ==========================
    for key in ["para1", "para2", "para3", "para4"]:
        txt = safe_block(chunks.get(key))
        if txt:
            doc.add_paragraph(txt)

    # ==========================
    # SIGNATURE
    # ==========================
    signature = safe_block(chunks.get("signature"))
    if signature:
        doc.add_paragraph(signature)

    # Sauvegarde
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)

    return output_path



# ============================================================================
# 6. Conversion DOCX -> PDF
# ============================================================================

def convert_html_to_pdf(html_content: str, pdf_path: str) -> str:
    """
    Convertit le HTML en PDF via xhtml2pdf.
    """
    pdf_path = os.path.abspath(pdf_path)
    os.makedirs(os.path.dirname(pdf_path) or ".", exist_ok=True)

    with open(pdf_path, "wb") as result_file:
        pisa.CreatePDF(html_content, dest=result_file)
    return pdf_path


# ============================================================================
# 7. Fonction high-level : Gemini + DOCX + PDF
# ============================================================================

def generate_personalized_cover_letter_docx_and_pdf(
    offer_parsed: Dict[str, Any],
    cv_parsed: Dict[str, Any],
    output_dir: str = ".",
    docx_filename: str = "lettre_motivation_gemini.docx",
    pdf_filename: Optional[str] = None,
    gender: str = "M",
    reference: Optional[str] = None,
    city_override: Optional[str] = None,
    date_override: Optional[str] = None,
    api_key: str = "",
    model_name: str = "gemini-1.5-flash"
) -> Dict[str, Any]:
    """
    Pipeline complet :
      1. Gemini génère une lettre structurée (JSON).
      2. On construit le DOCX.
      3. On convertit en PDF.

    Retourne:
      {
        "docx_path": "...",
        "pdf_path": "...",
        "chunks": {...}  # Contenu structuré pour le frontend
      }
    """
    os.makedirs(output_dir, exist_ok=True)
    docx_path = os.path.join(output_dir, docx_filename)

    # 1) Génération des chunks via Gemini
    chunks = generate_letter_structure_with_gemini(
        offer_parsed=offer_parsed,
        cv_parsed=cv_parsed,
        gender=gender,
        reference=reference,
        city_override=city_override,
        date_override=date_override,
        api_key=api_key,
        model_name=model_name
    )

    # 2) Génération du DOCX
    docx_path = build_cover_letter_docx_from_chunks(chunks, output_path=docx_path)

    # 3) Génération du PDF (via HTML pour éviter la dépendance Word)
    if pdf_filename is None:
        base, _ = os.path.splitext(docx_filename)
        pdf_filename = base + ".pdf"
    pdf_path = os.path.join(output_dir, pdf_filename)
    html_content = build_cover_letter_html_from_chunks(chunks)
    pdf_path = convert_html_to_pdf(html_content, pdf_path=pdf_path)

    return {
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "chunks": chunks,
    }


# ============================================================================
# 8. Démo rapide
# ============================================================================

if __name__ == "__main__":
    # Mini exemples fake, juste pour tester la structure.
    offer_demo = {
        "title": "Data Scientist Junior",
        "company_name": "Airbus",
        "location": "Toulouse",
        "contract_type": "CDI",
        "seniority_level": "Junior",
        "Education": ["Diplôme d'ingénieur en Data Science"],
        "hard_skills": ["Python", "SQL", "Machine Learning"],
        "soft_skills": ["travail en équipe", "communication"],
        "missions": [
            "Développer des modèles de machine learning en Python.",
            "Collaborer avec les équipes d'ingénierie."
        ],
        "requirements": [
            "1–2 ans d'expérience en data science (stages inclus)."
        ],
        "language": "fr"
    }

    cv_demo = {
        "first_name": "Theau",
        "last_name": "Aguet",
        "full_name": "Theau AGUET",
        "contacts": {
            "emails": ["theauaguetpro@gmail.com"],
            "phones": ["+33 7 82 01 83 97"],
            "locations": ["France"]
        },
        "skills": {
            "hard_skills": ["Python", "SQL", "Pandas", "Machine Learning"],
            "soft_skills": ["Travail en équipe", "Communication"],
            "languages": ["Français (C1)", "Anglais (C1)"]
        },
        "professional_experiences": [
            {
                "title": "Data Scientist Intern",
                "company": "Airbus",
                "location": "Toulouse",
                "start_date": "Feb 2024",
                "end_date": "Aug 2024",
                "description": "Travail sur des modèles de maintenance prédictive."
            }
        ],
        "education": [
            {
                "degree": "Diplôme d'ingénieur en Data Science",
                "school": "IMT Atlantique",
                "location": "France",
                "start_date": "2022",
                "end_date": "2025",
                "description": "Spécialisation en machine learning, statistiques et data engineering."
            }
        ],
    }

    print("[INFO] Génération de la lettre de motivation personnalisée (Gemini) en DOCX + PDF...")
    paths = generate_personalized_cover_letter_docx_and_pdf(
        offer_parsed=offer_demo,
        cv_parsed=cv_demo,
        output_dir=here("backend/generator/cover_letter"),
        docx_filename="lettre_motivation_gemini_demo.docx",
        gender="M",
        reference=None,
    )
    print("DOCX:", paths["docx_path"])
    print("PDF :", paths["pdf_path"])
